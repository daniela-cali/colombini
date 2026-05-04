"""
Genera il documento Word di presentazione funzionalità Colombini Gestionale
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ---------------------------------------------------------------------------
# Stili globali
# ---------------------------------------------------------------------------
style_normal = doc.styles['Normal']
style_normal.font.name = 'Calibri'
style_normal.font.size = Pt(11)

# Margini pagina
section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2)

# Colori aziendali
BLU    = RGBColor(0x1A, 0x56, 0x76)   # blu acqua profonda
ACQUA  = RGBColor(0x00, 0x9F, 0xBE)   # azzurro piscina
GRIGIO = RGBColor(0x4A, 0x4A, 0x4A)
VERDE  = RGBColor(0x1E, 0x8A, 0x44)
BIANCO = RGBColor(0xFF, 0xFF, 0xFF)

def set_cell_bg(cell, hex_color: str):
    """Imposta il colore di sfondo di una cella."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def set_table_borders(table, color='1A5676'):
    tbl  = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for side in ('top','left','bottom','right','insideH','insideV'):
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'),   'single')
        b.set(qn('w:sz'),    '4')
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), color)
        tblBorders.append(b)
    tblPr.append(tblBorders)

def heading1(text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text.upper())
    run.font.name  = 'Calibri'
    run.font.size  = Pt(15)
    run.font.bold  = True
    run.font.color.rgb = BLU
    # linea sotto
    pPr   = p._p.get_or_add_pPr()
    pBdr  = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '009FBE')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def heading2(text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(text)
    run.font.name  = 'Calibri'
    run.font.size  = Pt(12)
    run.font.bold  = True
    run.font.color.rgb = ACQUA
    return p

def heading3(text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.font.name  = 'Calibri'
    run.font.size  = Pt(11)
    run.font.bold  = True
    run.font.color.rgb = GRIGIO
    return p

def body(text: str, indent: float = 0):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(indent)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run.font.color.rgb = GRIGIO
    return p

def bullet(text: str, indent: float = 0.5):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent   = Cm(indent)
    p.paragraph_format.space_after   = Pt(2)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run.font.color.rgb = GRIGIO
    return p

def note_box(text: str):
    """Paragrafo con sfondo azzurro chiaro come box nota."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    run.font.name   = 'Calibri'
    run.font.size   = Pt(10.5)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x00, 0x5F, 0x7A)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  'D6F0F8')
    pPr.append(shd)
    return p

def page_break():
    doc.add_page_break()

# ===========================================================================
# COPERTINA
# ===========================================================================

# Spazio bianco in alto
for _ in range(4):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('COLOMBINI S.n.c.')
run.font.name  = 'Calibri'
run.font.size  = Pt(28)
run.font.bold  = True
run.font.color.rgb = BLU

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Gestionale Interventi e Portale Cliente')
run.font.name  = 'Calibri'
run.font.size  = Pt(18)
run.font.color.rgb = ACQUA

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Documento di Presentazione Funzionalità')
run.font.name  = 'Calibri'
run.font.size  = Pt(14)
run.font.bold  = True
run.font.color.rgb = GRIGIO

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Funzionalità Implementate e Roadmap di Sviluppo')
run.font.name   = 'Calibri'
run.font.size   = Pt(12)
run.font.italic = True
run.font.color.rgb = GRIGIO

for _ in range(3):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(f'Versione 1.0  —  {datetime.date.today().strftime("%B %Y")}')
run.font.name  = 'Calibri'
run.font.size  = Pt(11)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

page_break()

# ===========================================================================
# INDICE (manuale)
# ===========================================================================
heading1('Indice')

indice = [
    ('1.', 'Panoramica del Sistema',                  '3'),
    ('2.', 'Architettura Tecnologica',                '3'),
    ('3.', 'Gestione Utenti e Accessi',               '4'),
    ('4.', 'Moduli Implementati',                     '5'),
    ('  4.1', 'Dashboard',                            '5'),
    ('  4.2', 'Anagrafica Clienti',                   '6'),
    ('  4.3', 'Gestione Interventi',                  '7'),
    ('  4.4', 'Gestione Tecnici',                     '8'),
    ('  4.5', 'Portale Cliente',                      '9'),
    ('  4.6', 'Impostazioni',                         '10'),
    ('5.', 'Roadmap: Funzionalità in Sviluppo',       '11'),
    ('  5.1', 'Gestione Impianti',                    '11'),
    ('  5.2', 'Gestione Preventivi',                  '12'),
    ('  5.3', 'Catalogo Prodotti e Ricambi',          '13'),
    ('  5.4', 'Report e Analisi',                     '13'),
    ('  5.5', 'Profilo Utente e Notifiche',           '14'),
    ('  5.6', 'Miglioramenti Portale Cliente',        '14'),
    ('6.', 'Riepilogo Stato Sviluppo',                '15'),
]

tbl = doc.add_table(rows=len(indice), cols=3)
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, (num, titolo, pag) in enumerate(indice):
    tbl.rows[i].cells[0].paragraphs[0].add_run(num).font.size = Pt(11)
    tbl.rows[i].cells[1].paragraphs[0].add_run(titolo).font.size = Pt(11)
    r = tbl.rows[i].cells[2].paragraphs[0]
    r.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r.add_run(pag).font.size = Pt(11)
    tbl.rows[i].cells[0].width = Cm(1.5)
    tbl.rows[i].cells[2].width = Cm(1.5)

page_break()

# ===========================================================================
# 1. PANORAMICA
# ===========================================================================
heading1('1. Panoramica del Sistema')

body(
    'Il Gestionale Colombini è un\'applicazione web sviluppata su misura per Colombini S.n.c., '
    'azienda specializzata nella manutenzione e installazione di piscine, addolcitori '
    'e acquedotti. Il sistema centralizza la gestione operativa e '
    'offre ai clienti un portale dedicato per seguire le proprie richieste di assistenza.'
)

doc.add_paragraph()
body('Le finalità principali del sistema sono:', 0)
bullet('Centralizzare l\'anagrafica dei clienti con ricerca avanzata e importazione massiva da CSV')
bullet('Pianificare e tracciare gli interventi tecnici, dai sopralluoghi alle manutenzioni')
bullet('Fornire ai tecnici un\'agenda digitale aggiornata in tempo reale')
bullet('Permettere ai clienti di aprire richieste di assistenza e monitorarne lo stato')
bullet('Supportare la crescita aziendale con moduli espandibili per preventivi, impianti e report')

# ===========================================================================
# 2. ARCHITETTURA
# ===========================================================================
heading1('2. Architettura Tecnologica')

body('Il sistema è costruito su uno stack moderno, stabile e orientato alla sicurezza.')

tbl = doc.add_table(rows=9, cols=2)
tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
set_table_borders(tbl)

headers = ['Componente', 'Tecnologia']
for j, h in enumerate(headers):
    c = tbl.rows[0].cells[j]
    set_cell_bg(c, '1A5676')
    r = c.paragraphs[0].add_run(h)
    r.font.bold  = True
    r.font.color.rgb = BIANCO
    r.font.size  = Pt(11)

righe = [
    ('Framework Backend',    'PHP 8.2+ · CodeIgniter 4.7'),
    ('Autenticazione',       'CodeIgniter Shield 1.3 (sessioni + gruppi/ruoli)'),
    ('Database',             'MySQL 8.0, charset utf8mb4'),
    ('ORM / Query Builder',  'CodeIgniter Model layer (migrations incluse)'),
    ('Geocodifica',          'OpenStreetMap Nominatim API (gratuita)'),
    ('Import dati',          'Parser CSV nativo PHP con auto-detection separatore'),
    ('Frontend',             'Bootstrap 5, HTML5, JavaScript vanilla'),
    ('Hosting previsto',     'Server LAMP (Linux, Apache, MySQL, PHP)'),
]

for i, (comp, tech) in enumerate(righe, start=1):
    tbl.rows[i].cells[0].paragraphs[0].add_run(comp).font.size = Pt(11)
    tbl.rows[i].cells[1].paragraphs[0].add_run(tech).font.size = Pt(11)
    if i % 2 == 0:
        set_cell_bg(tbl.rows[i].cells[0], 'EBF7FB')
        set_cell_bg(tbl.rows[i].cells[1], 'EBF7FB')

page_break()

# ===========================================================================
# 3. GESTIONE UTENTI E ACCESSI
# ===========================================================================
heading1('3. Gestione Utenti e Accessi')

body(
    'Il sistema implementa un modello RBAC (Role-Based Access Control) con quattro ruoli distinti, '
    'ciascuno con accesso a sezioni e funzionalità specifiche.'
)

doc.add_paragraph()

tbl = doc.add_table(rows=5, cols=4)
tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
set_table_borders(tbl)
tbl.columns[0].width = Cm(3)
tbl.columns[1].width = Cm(3)
tbl.columns[2].width = Cm(5)
tbl.columns[3].width = Cm(4.5)

hdrs = ['Ruolo', 'Area di accesso', 'Funzionalità principali', 'Creato da']
for j, h in enumerate(hdrs):
    c = tbl.rows[0].cells[j]
    set_cell_bg(c, '1A5676')
    r = c.paragraphs[0].add_run(h)
    r.font.bold = True
    r.font.color.rgb = BIANCO
    r.font.size = Pt(11)

ruoli = [
    ('Amministratore', 'Gestionale completo',
     'Accesso a tutti i moduli, impostazioni, gestione utenti',
     'Configurazione iniziale'),
    ('Staff',          'Gestionale completo',
     'Clienti, interventi, tecnici (senza impostazioni avanzate)',
     'Amministratore'),
    ('Tecnico',        'Gestionale (vista ridotta)',
     'Dashboard agenda, dettaglio interventi assegnati',
     'Amministratore'),
    ('Cliente',        'Portale Cliente',
     'Nuova richiesta di assistenza, stato richieste',
     'Admin/Staff da anagrafica cliente'),
]

for i, (ruolo, area, funz, crea) in enumerate(ruoli, start=1):
    row = tbl.rows[i]
    row.cells[0].paragraphs[0].add_run(ruolo).font.size = Pt(11)
    row.cells[1].paragraphs[0].add_run(area).font.size = Pt(11)
    row.cells[2].paragraphs[0].add_run(funz).font.size = Pt(11)
    row.cells[3].paragraphs[0].add_run(crea).font.size = Pt(11)
    if i % 2 == 0:
        for j in range(4):
            set_cell_bg(row.cells[j], 'EBF7FB')

doc.add_paragraph()
note_box(
    '★  Il login avviene tramite username e password. I clienti vengono reindirizzati '
    'automaticamente al portale dedicato, mentre staff, tecnici e admin accedono al '
    'gestionale interno.'
)

page_break()

# ===========================================================================
# 4. MODULI IMPLEMENTATI
# ===========================================================================
heading1('4. Moduli Implementati')

# 4.1 Dashboard
heading2('4.1  Dashboard')
body(
    'La dashboard è differenziata per ruolo e fornisce una vista immediata sullo stato '
    'operativo dell\'azienda.'
)
doc.add_paragraph()
heading3('Dashboard Amministratore / Staff')
bullet('Contatore clienti totali attivi')
bullet('Numero impianti registrati')
bullet('Interventi aperti (pianificati + in corso)')
bullet('Interventi completati nel mese corrente')
bullet('Accesso rapido alle sezioni più utilizzate')

heading3('Dashboard Tecnico')
bullet('Agenda interventi assegnati al tecnico loggato')
bullet('Filtro per data: oggi, settimana, mese')
bullet('Stato di ogni intervento con badge colorato (pianificato / in corso / completato / annullato)')
bullet('Riepilogo numerico: pianificati, in corso, completati del mese')

doc.add_paragraph()
note_box(
    'I tecnici accedono esclusivamente alla propria agenda e ai dettagli degli interventi '
    'loro assegnati. Non hanno visibilità sulle anagrafiche o sulle configurazioni di sistema.'
)

# 4.2 Clienti
heading2('4.2  Anagrafica Clienti')
body(
    'Il modulo clienti è il cuore dell\'anagrafica aziendale. Supporta sia privati che società, '
    'con un sistema di codici interni univoci generati automaticamente (es. INT-0001).'
)

doc.add_paragraph()
heading3('Funzionalità CRUD')
bullet('Creazione cliente: tipo (persona fisica / società), dati anagrafici, indirizzo, recapiti, note')
bullet('Codice interno generato automaticamente in sequenza')
bullet('Modifica e cancellazione con verifica integrità')
bullet('Vista dettaglio con riepilogo interventi e richieste collegate')

heading3('Ricerca Avanzata')
bullet('Ricerca testuale su: ragione sociale, cognome, nome, codice interno, P.IVA, codice fiscale')
bullet('Risultati in tempo reale con paginazione')

heading3('Importazione Massiva da CSV')
bullet('Upload di file CSV esportati da altri gestionali')
bullet('Auto-detection del separatore (; o ,) e dell\'encoding (UTF-8 / ISO-8859-1)')
bullet('Interfaccia di mapping colonne: l\'utente associa le colonne del CSV ai campi del sistema')
bullet('Upsert intelligente: aggiorna i clienti già presenti (per codice), inserisce i nuovi')
bullet('Report a fine import: inserimenti, aggiornamenti, errori')

heading3('Geocodifica Indirizzi')
bullet('Conversione automatica degli indirizzi in coordinate GPS (lat/lng)')
bullet('Tramite OpenStreetMap Nominatim API (nessun costo)')
bullet('Coordinate salvate per utilizzo futuro su mappe')

heading3('Accesso Portale Cliente')
bullet('Creazione credenziali portale direttamente dalla scheda cliente')
bullet('Username e password assegnati per persona fisica o responsabile aziendale')
bullet('Il cliente può accedere immediatamente al portale con le credenziali create')

# 4.3 Interventi
heading2('4.3  Gestione Interventi')
body(
    'Il modulo interventi traccia l\'intero ciclo di vita di ogni visita tecnica, '
    'dalla pianificazione alla chiusura.'
)

doc.add_paragraph()
heading3('Dati dell\'Intervento')
bullet('Tipo intervento: Piscine · Addolcitori · Acquedotti · Commerciale')
bullet('Cliente e luogo di esecuzione')
bullet('Tecnico assegnato')
bullet('Data e ora pianificata')
bullet('Durata del viaggio e durata effettiva dell\'intervento')
bullet('Data di completamento (impostata automaticamente alla chiusura)')
bullet('Descrizione del lavoro svolto e note interne')
bullet('Collegamento opzionale alla richiesta di assistenza di origine')

heading3('Ciclo di Vita (Stati)')
tbl = doc.add_table(rows=5, cols=3)
set_table_borders(tbl, '009FBE')
for j, h in enumerate(['Stato', 'Significato', 'Transizioni possibili']):
    c = tbl.rows[0].cells[j]
    set_cell_bg(c, '009FBE')
    r = c.paragraphs[0].add_run(h)
    r.font.bold = True
    r.font.color.rgb = BIANCO
    r.font.size = Pt(11)
stati = [
    ('Pianificato',  'Intervento schedulato, non ancora avviato',      '→ In corso, Annullato'),
    ('In corso',     'Tecnico sul posto / lavori in esecuzione',        '→ Completato, Annullato'),
    ('Completato',   'Lavoro terminato, data chiusura registrata',      '— (stato finale)'),
    ('Annullato',    'Intervento cancellato',                           '— (stato finale)'),
]
for i, (s, sig, tr) in enumerate(stati, 1):
    tbl.rows[i].cells[0].paragraphs[0].add_run(s).font.size = Pt(11)
    tbl.rows[i].cells[1].paragraphs[0].add_run(sig).font.size = Pt(11)
    tbl.rows[i].cells[2].paragraphs[0].add_run(tr).font.size = Pt(11)
    if i % 2 == 0:
        for j in range(3): set_cell_bg(tbl.rows[i].cells[j], 'EBF7FB')

doc.add_paragraph()
heading3('Funzionalità Operative')
bullet('Assegnazione tecnico con un click dalla lista o dalla scheda intervento')
bullet('Chiusura rapida con registrazione automatica della data')
bullet('Filtri per stato, tipo e tecnico nella lista interventi')
bullet('Generazione PDF intervento (predisposta, in completamento)')

# 4.4 Tecnici
heading2('4.4  Gestione Tecnici')
body(
    'Il modulo tecnici gestisce l\'organico operativo e i relativi orari di lavoro.'
)
doc.add_paragraph()
heading3('Anagrafica Tecnico')
bullet('Nome, cognome, username di accesso, telefono')
bullet('Ruolo nel sistema (tecnico, staff, admin)')
bullet('Lista interventi assegnati con stati')

heading3('Orari di Lavoro Settimanali')
bullet('Configurazione per ogni giorno della settimana (lunedì–domenica)')
bullet('Orario di inizio e fine turno')
bullet('Fascia di pausa pranzo (inizio e fine)')
bullet('Attivazione/disattivazione giornata lavorativa')
bullet('Gli orari default aziendali sono configurabili nelle Impostazioni')

heading3('Statistiche Tecnico')
bullet('Numero interventi pianificati, in corso, completati del mese')
bullet('Vista riassuntiva per il responsabile')

# 4.5 Portale Cliente
heading2('4.5  Portale Cliente')
body(
    'Il portale è un\'area separata, accessibile ai soli clienti registrati. '
    'Offre un\'interfaccia semplice e intuitiva per comunicare con l\'azienda.'
)
doc.add_paragraph()
heading3('Dashboard Portale')
bullet('Lista di tutte le richieste di assistenza inviate dal cliente')
bullet('Stato di ogni richiesta con badge colorato (Nuova / In lavorazione / Chiusa)')
bullet('Accesso al dettaglio di ogni richiesta')

heading3('Nuova Richiesta di Assistenza')
bullet('Selezione tipo intervento richiesto: Piscine, Addolcitori, Acquedotti, Commerciale')
bullet('Indicazione del luogo di intervento (diverso dalla sede se necessario)')
bullet('Nome del referente e telefono di contatto')
bullet('Descrizione del problema o della richiesta')

heading3('Dettaglio Richiesta')
bullet('Tutti i dati inseriti alla creazione')
bullet('Stato corrente della lavorazione')
bullet('Data di inserimento')

doc.add_paragraph()
note_box(
    'Le richieste del portale vengono ricevute dallo staff nel gestionale interno, '
    'che può associarle a un intervento e aggiornarne lo stato. Il cliente vede '
    'immediatamente il cambio di stato nella propria area.'
)

# 4.6 Impostazioni
heading2('4.6  Impostazioni')
body(
    'L\'area impostazioni è riservata agli amministratori e permette di personalizzare '
    'il sistema in base alle esigenze aziendali.'
)
doc.add_paragraph()
heading3('Parametri Aziendali')
bullet('Nome azienda, indirizzo sede, coordinate GPS della sede')
bullet('Questi dati sono usati come default per geocodifica e futura cartografia')

heading3('Orari Default Tecnici')
bullet('Orario di inizio e fine lavoro standard')
bullet('Applicato automaticamente ai nuovi tecnici creati')

heading3('Durate Standard Interventi')
bullet('Durata predefinita per ogni tipo di intervento (Piscine: 120 min, Addolcitori: 60 min, ecc.)')
bullet('Modificabile per adattarsi ai tempi reali dell\'azienda')

heading3('Gestione Utenti Portale')
bullet('Lista di tutti i clienti con accesso al portale')
bullet('Creazione nuovo account portale (associato a cliente esistente)')
bullet('Modifica e revoca accessi')

heading3('Gestione Utenti App')
bullet('Lista di admin, staff e tecnici')
bullet('Creazione nuovi utenti con assegnazione ruolo')
bullet('Modifica dati, cambio password, eliminazione')

page_break()

# ===========================================================================
# 5. ROADMAP
# ===========================================================================
heading1('5. Roadmap: Funzionalità in Sviluppo')

body(
    'I moduli descritti in questa sezione sono già predisposti nella struttura dell\'applicazione '
    '(routing, controller placeholder) e costituiscono la naturale evoluzione del sistema. '
    'Lo sviluppo procederà per priorità concordate.'
)

# 5.1 Impianti
heading2('5.1  Gestione Impianti  [PRIORITÀ ALTA]')
body(
    'Ogni cliente può avere uno o più impianti registrati (piscine, addolcitori, '
    'acquedotti, impianti commerciali). La scheda impianto diventa il punto di riferimento '
    'per tutta la storia manutentiva.'
)
doc.add_paragraph()
heading3('Scheda Impianto')
bullet('Tipo impianto (piscina / addolcitore / acquedotto / commerciale)')
bullet('Cliente di appartenenza')
bullet('Indirizzo di installazione (può differire dalla sede del cliente)')
bullet('Anno di installazione e dati tecnici (marca, modello, matricola)')
bullet('Foto dell\'impianto')
bullet('Note tecniche di riferimento')

heading3('Storico Interventi per Impianto')
bullet('Tutti gli interventi eseguiti sull\'impianto, in ordine cronologico')
bullet('Prodotti/ricambi utilizzati in ogni intervento')
bullet('Costo storico delle manutenzioni')

heading3('Scadenze e Manutenzioni Programmate')
bullet('Registro scadenze per tipo: manutenzione ordinaria, trattamento acque, revisione filtri')
bullet('Alert automatici allo staff quando si avvicina una scadenza')
bullet('Pianificazione preventiva degli interventi di routine')

heading3('Integrazione con Moduli Esistenti')
bullet('Collegamento interventi → impianto (invece del solo cliente)')
bullet('Filtro clienti per impianti in scadenza')

# 5.2 Preventivi
heading2('5.2  Gestione Preventivi  [PRIORITÀ ALTA]')
body(
    'Il modulo preventivi permette di creare offerte commerciali strutturate, '
    'inviarle ai clienti e convertirle in ordini/interventi con un click.'
)
doc.add_paragraph()
heading3('Creazione Preventivo')
bullet('Intestazione cliente, data emissione, validità offerta')
bullet('Righe di dettaglio: prodotti/servizi dal catalogo, quantità, prezzo unitario, sconto')
bullet('IVA per riga o globale')
bullet('Totale imponibile, IVA e totale documento')
bullet('Note e condizioni di pagamento')
bullet('Numerazione automatica progressiva')

heading3('Ciclo di Vita del Preventivo')
bullet('Bozza → Inviato → Accettato / Rifiutato / Scaduto')
bullet('Conversione accettato → Intervento (o serie di interventi)')
bullet('Storico revisioni del preventivo')

heading3('Stampa e Invio')
bullet('Generazione PDF professionale con logo e intestazione azienda')
bullet('Invio via email direttamente dal gestionale')
bullet('Download PDF per invio manuale')

# 5.3 Prodotti
heading2('5.3  Catalogo Prodotti e Ricambi  [PRIORITÀ MEDIA]')
body(
    'Un catalogo centralizzato di prodotti e ricambi usati negli interventi, '
    'utile per la compilazione rapida di preventivi e per il tracciamento dei materiali.'
)
doc.add_paragraph()
bullet('Scheda prodotto: codice, descrizione, categoria, fornitore, prezzo di acquisto, prezzo di vendita')
bullet('Unità di misura (pezzo, litro, kg, mt…)')
bullet('Gestione giacenza di magazzino (opzionale)')
bullet('Categorie personalizzabili: chimici, filtri, pompe, accessori, manodopera')
bullet('Collegamento ai preventivi e alle schede intervento')
bullet('Storico utilizzo per prodotto')

# 5.4 Report
heading2('5.4  Report e Analisi  [PRIORITÀ MEDIA]')
body(
    'Il modulo report trasforma i dati operativi in informazioni utili per '
    'le decisioni aziendali.'
)
doc.add_paragraph()
heading3('Report Operativi')
bullet('Interventi per periodo: totali, per tipo, per tecnico, per stato')
bullet('Produttività tecnici: ore lavorate, interventi completati, tempi medi')
bullet('Clienti più attivi: numero interventi, valore stimato')
bullet('Richieste portale: volumi, tempi di risposta, tassi di chiusura')

heading3('Report Economici (con Preventivi)')
bullet('Fatturato preventivato vs consuntivato per periodo')
bullet('Tasso di conversione preventivi')
bullet('Valore medio per tipo di intervento')

heading3('Esportazione')
bullet('Export CSV per elaborazione in Excel')
bullet('Export PDF stampabile')
bullet('Grafici visuali (istogrammi, torte) su browser')

# 5.5 Profilo e Notifiche
heading2('5.5  Profilo Utente e Notifiche  [PRIORITÀ BASSA]')
doc.add_paragraph()
heading3('Gestione Profilo')
bullet('Modifica dati personali (nome, cognome, telefono)')
bullet('Cambio password in autonomia')
bullet('Preferenze di visualizzazione (lingua, timezone)')

heading3('Sistema di Notifiche')
bullet('Notifiche in-app per nuove richieste portale (per staff)')
bullet('Notifiche per interventi assegnati (per tecnici)')
bullet('Invio email opzionale sulle principali notifiche')
bullet('Centro notifiche con storico letture')

# 5.6 Miglioramenti Portale
heading2('5.6  Miglioramenti Portale Cliente  [PRIORITÀ MEDIA]')
doc.add_paragraph()
heading3('Vista Impianti del Cliente')
bullet('Il cliente vede i propri impianti registrati e le scadenze')
bullet('Storico interventi eseguiti sui propri impianti')

heading3('Documenti e Preventivi')
bullet('Download dei preventivi accettati in formato PDF')
bullet('Archivio documenti (verbali, certificati di manutenzione)')

heading3('Comunicazione Bidirezionale')
bullet('Risposta dello staff alle richieste direttamente nel thread della richiesta')
bullet('Notifica email al cliente ad ogni aggiornamento')

page_break()

# ===========================================================================
# 6. RIEPILOGO STATO SVILUPPO
# ===========================================================================
heading1('6. Riepilogo Stato di Sviluppo')

body('La tabella seguente mostra lo stato di tutti i moduli del sistema.')
doc.add_paragraph()

moduli = [
    ('Autenticazione e RBAC',           '✔ Completato',     'Produzione'),
    ('Dashboard (admin e tecnico)',      '✔ Completato',     'Produzione'),
    ('Anagrafica Clienti',              '✔ Completato',     'Produzione'),
    ('Import CSV Clienti',              '✔ Completato',     'Produzione'),
    ('Geocodifica Indirizzi',           '✔ Completato',     'Produzione'),
    ('Portale Cliente',                 '✔ Completato',     'Produzione'),
    ('Gestione Interventi',             '✔ Completato',     'Produzione'),
    ('Gestione Tecnici + Orari',        '✔ Completato',     'Produzione'),
    ('Impostazioni Sistema',            '✔ Completato',     'Produzione'),
    ('PDF Intervento',                  '◑ In completamento', 'Q3 2025'),
    ('Gestione Impianti',               '○ Da sviluppare',  'Q3 2025'),
    ('Gestione Preventivi',             '○ Da sviluppare',  'Q3-Q4 2025'),
    ('Catalogo Prodotti/Ricambi',       '○ Da sviluppare',  'Q4 2025'),
    ('Report e Analisi',                '○ Da sviluppare',  'Q4 2025'),
    ('Notifiche in-app / Email',        '○ Da sviluppare',  'Q1 2026'),
    ('Profilo Utente Avanzato',         '○ Da sviluppare',  'Q1 2026'),
    ('Miglioramenti Portale Cliente',   '○ Da sviluppare',  'Q1 2026'),
]

tbl = doc.add_table(rows=len(moduli)+1, cols=3)
set_table_borders(tbl)
tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

for j, h in enumerate(['Modulo', 'Stato', 'Rilascio Previsto']):
    c = tbl.rows[0].cells[j]
    set_cell_bg(c, '1A5676')
    r = c.paragraphs[0].add_run(h)
    r.font.bold = True
    r.font.color.rgb = BIANCO
    r.font.size = Pt(11)

for i, (mod, stato, rilascio) in enumerate(moduli, 1):
    row = tbl.rows[i]
    row.cells[0].paragraphs[0].add_run(mod).font.size = Pt(11)
    s_run = row.cells[1].paragraphs[0].add_run(stato)
    s_run.font.size = Pt(11)
    if '✔' in stato:
        s_run.font.color.rgb = VERDE
        s_run.font.bold = True
    elif '◑' in stato:
        s_run.font.color.rgb = RGBColor(0xD4, 0x7C, 0x00)
        s_run.font.bold = True
    else:
        s_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    row.cells[2].paragraphs[0].add_run(rilascio).font.size = Pt(11)
    if i % 2 == 0:
        for j in range(3): set_cell_bg(row.cells[j], 'EBF7FB')

doc.add_paragraph()
note_box(
    'Le date di rilascio sono indicative e soggette a variazioni in base alle priorità '
    'concordate. I moduli ✔ Completato sono pienamente operativi nella versione demo e '
    'in produzione.'
)

# Piè di pagina
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(
    f'Colombini S.n.c.  —  Documento riservato  —  '
    f'Generato il {datetime.date.today().strftime("%d/%m/%Y")}'
)
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
run.font.italic = True

# ---------------------------------------------------------------------------
# SALVATAGGIO
# ---------------------------------------------------------------------------
out_path = r'd:\Programmazione\colombini\Colombini_Gestionale_Funzionalita.docx'
doc.save(out_path)
print(f'Documento salvato: {out_path}')
